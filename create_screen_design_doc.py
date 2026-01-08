"""
nCompliance 화면 구성안 문서 생성 스크립트
캡처된 스크린샷을 포함한 Word 문서를 생성합니다.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from datetime import datetime

# 화면 정보 정의
SCREENS = [
    {
        "id": "SCR-001",
        "name": "로그인",
        "image": "01_login.png",
        "description": "시스템 로그인 화면",
        "features": [
            "아이디/비밀번호 입력 폼",
            "로그인 버튼",
            "시스템 브랜드 로고 표시",
            "배경 그라데이션 디자인"
        ],
        "access": "전체 (비로그인)"
    },
    {
        "id": "SCR-002",
        "name": "Home (사규 브라우저)",
        "image": "02_home.png",
        "description": "메인 대시보드 화면으로 사규를 트리 구조로 탐색할 수 있습니다.",
        "features": [
            "카테고리별 탭 (전체/정책·방침/규정/지침) 및 사규 수 표시",
            "그룹별 트리 구조 (이사회, HR, 재무, 구매 등)",
            "즐겨찾기 그룹 (상단 고정)",
            "사규 선택 시 오른쪽 패널에 본문 표시",
            "사규 바로가기 링크 (외부 URL 이동)"
        ],
        "access": "전체 (로그인 필요)"
    },
    {
        "id": "SCR-003",
        "name": "사규 목록",
        "image": "03_regulation_list.png",
        "description": "등록된 전체 사규 목록을 조회하고 검색할 수 있습니다.",
        "features": [
            "검색 필터 (검색어, 유형, 상태, 의무준수, 임직원 공개, 그룹, 담당부서, 담당자)",
            "검색 결과 테이블 (사규코드, 그룹, 사규명, 유형, 의무, 담당부서, 담당자, 상태, 공개, 시행일, 링크)",
            "Excel 다운로드 기능",
            "사규 등록 버튼",
            "페이징 처리"
        ],
        "access": "책임부서 담당자 이상"
    },
    {
        "id": "SCR-004",
        "name": "사규 등록/수정",
        "image": "04_regulation_create.png",
        "description": "새로운 사규를 등록하거나 기존 사규를 수정합니다.",
        "features": [
            "기본 정보 입력 (사규명, 분류, 그룹, 설명)",
            "관리 정보 입력 (책임부서, 담당자, 책임자)",
            "옵션 설정 (의무준수, 임직원 공개)",
            "날짜 정보 (시행일, 정기검토예정일)",
            "관련 사규 연결 (상위 사규, 관련 사규)",
            "태그 입력",
            "접근 권한 설정",
            "파일 업로드 (원본 파일)",
            "규정 본문 입력",
            "참조 링크 입력"
        ],
        "access": "책임부서 담당자 이상"
    },
    {
        "id": "SCR-005",
        "name": "사규 상세",
        "image": "04_regulation_detail.png",
        "description": "사규의 상세 정보를 확인하고 버전 이력을 조회합니다.",
        "features": [
            "기본 정보 표시 (코드, 명칭, 분류, 그룹, 버전, 상태)",
            "관리 정보 표시 (책임부서, 담당자, 책임자)",
            "날짜 정보 표시 (시행일, 정기검토예정일, 등록일)",
            "버전 이력 목록",
            "관련 사규 링크",
            "태그 목록",
            "원본 파일 다운로드",
            "수정/삭제 버튼 (권한에 따라)"
        ],
        "access": "전체 (로그인 필요, 권한에 따른 데이터 필터링)"
    },
    {
        "id": "SCR-006",
        "name": "태그 관리",
        "image": "05_tag_list.png",
        "description": "사규에 연결된 태그를 카테고리별로 그룹핑하여 표시합니다.",
        "features": [
            "전체 태그 수 및 사규 수 표시",
            "8개 카테고리 그룹핑 (이사회·지배구조, 인사·HR, 재무·회계, 컴플라이언스·준법, 보안·정보보호, 구매·조달, 감사·내부감사, 기타)",
            "태그별 연결된 사규 수 표시",
            "태그 크기 동적 조절 (사규 수에 따라)",
            "태그 클릭 시 해당 사규 목록으로 이동",
            "태그 검색 기능"
        ],
        "access": "책임부서 담당자 이상"
    },
    {
        "id": "SCR-007",
        "name": "제개정 이력",
        "image": "06_change_history.png",
        "description": "사규의 제정/개정/폐지 이력을 조회합니다.",
        "features": [
            "기간별 검색 (시작일, 종료일)",
            "변경유형 필터 (전체/제정/개정/폐지)",
            "이력 테이블 (사규코드, 그룹, 분류, 사규명, 버전, 변경사유, 작성자, 등록일)",
            "Excel 다운로드 기능",
            "페이징 처리"
        ],
        "access": "책임부서 담당자 이상 (담당자는 소관 사규만)"
    },
    {
        "id": "SCR-008",
        "name": "만료예정",
        "image": "07_expiry_report.png",
        "description": "정기검토예정일이 도래하는 사규 목록을 조회합니다.",
        "features": [
            "조회 기간 선택 (7일, 14일, 30일, 60일, 90일 이내)",
            "만료예정 사규 테이블 (사규코드, 사규명, 분류, 책임부서, 정기검토예정일, 남은일수)",
            "남은 일수에 따른 색상 표시 (위험/경고/주의)",
            "Excel 다운로드 기능"
        ],
        "access": "책임부서 담당자 이상 (담당자는 소관 사규만)"
    },
    {
        "id": "SCR-009",
        "name": "알림 목록",
        "image": "08_notification_list.png",
        "description": "수신한 알림 목록을 확인하고 관리합니다.",
        "features": [
            "알림 목록 (유형, 제목, 내용, 수신일, 읽음여부)",
            "읽음/읽지않음 상태 표시",
            "알림 클릭 시 상세 내용 및 관련 사규 이동",
            "전체 읽음 처리 버튼",
            "알림 발송 버튼 (관리자/준법지원인만)"
        ],
        "access": "전체 (로그인 필요)"
    },
    {
        "id": "SCR-010",
        "name": "알림 발송",
        "image": "09_notification_create.png",
        "description": "시스템 관리자 또는 준법지원인이 알림을 발송합니다.",
        "features": [
            "알림 유형 선택 (제개정알림, 만료예정알림, 검토요청알림, 시스템알림)",
            "제목 및 내용 입력",
            "발송 대상 선택 (전체 사용자, 역할별, 개별 사용자)",
            "관련 사규 선택 (선택사항)",
            "발송 버튼"
        ],
        "access": "준법지원인, 시스템 관리자"
    },
    {
        "id": "SCR-011",
        "name": "코드 관리",
        "image": "10_code_list.png",
        "description": "시스템에서 사용하는 공통코드를 관리합니다.",
        "features": [
            "코드 유형별 필터링",
            "코드 목록 테이블 (코드유형, 코드, 코드명, 설명, 정렬순서, 사용여부, 시스템코드)",
            "코드 등록/수정/삭제 기능",
            "시스템 코드는 삭제 불가"
        ],
        "access": "시스템 관리자"
    },
]

def set_cell_shading(cell, fill_color):
    """셀 배경색 설정"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def create_document():
    doc = Document()
    
    # 문서 스타일 설정
    style = doc.styles['Normal']
    font = style.font
    font.name = '맑은 고딕'
    font.size = Pt(10)
    
    # 제목 스타일 설정
    for i in range(1, 4):
        heading_style = doc.styles[f'Heading {i}']
        heading_style.font.name = '맑은 고딕'
        heading_style.font.bold = True
    
    # ========== 표지 ==========
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    title = doc.add_paragraph()
    title_run = title.add_run("nCompliance")
    title_run.bold = True
    title_run.font.size = Pt(36)
    title_run.font.name = '맑은 고딕'
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run("사규관리 시스템")
    subtitle_run.font.size = Pt(24)
    subtitle_run.font.name = '맑은 고딕'
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    doc_title = doc.add_paragraph()
    doc_title_run = doc_title.add_run("화면 구성안")
    doc_title_run.bold = True
    doc_title_run.font.size = Pt(28)
    doc_title_run.font.name = '맑은 고딕'
    doc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # 문서 정보 테이블
    info_table = doc.add_table(rows=4, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    info_data = [
        ("문서 버전", "1.0"),
        ("작성일", datetime.now().strftime("%Y년 %m월 %d일")),
        ("프로젝트", "nCompliance (Corporate Regulation Management System)"),
        ("작성자", "개발팀"),
    ]
    
    for i, (label, value) in enumerate(info_data):
        info_table.cell(i, 0).text = label
        info_table.cell(i, 1).text = value
        info_table.cell(i, 0).paragraphs[0].runs[0].bold = True
    
    # 페이지 나누기
    doc.add_page_break()
    
    # ========== 목차 ==========
    doc.add_heading("목차", level=1)
    doc.add_paragraph()
    
    toc = doc.add_paragraph()
    toc.add_run("1. 개요").bold = True
    doc.add_paragraph("2. 화면 목록", style='List Number')
    
    for i, screen in enumerate(SCREENS, start=1):
        doc.add_paragraph(f"   {i}. {screen['name']} ({screen['id']})")
    
    doc.add_page_break()
    
    # ========== 1. 개요 ==========
    doc.add_heading("1. 개요", level=1)
    
    doc.add_paragraph(
        "본 문서는 nCompliance(사규관리 시스템)의 화면 구성안을 정의합니다. "
        "각 화면의 레이아웃, 구성 요소, 주요 기능을 설명하며, 실제 구현된 화면 캡처를 포함합니다."
    )
    
    doc.add_paragraph()
    doc.add_heading("1.1 시스템 개요", level=2)
    doc.add_paragraph(
        "nCompliance는 기업의 사규(정책, 규정, 지침, 매뉴얼 등)를 체계적으로 관리하고, "
        "제·개정·폐지 절차를 표준화하며, 임직원에게 사규 열람 서비스를 제공하는 웹 기반 시스템입니다."
    )
    
    doc.add_paragraph()
    doc.add_heading("1.2 화면 구성 요약", level=2)
    
    # 화면 목록 테이블
    summary_table = doc.add_table(rows=len(SCREENS)+1, cols=4)
    summary_table.style = 'Table Grid'
    
    headers = ["화면 ID", "화면명", "설명", "접근 권한"]
    header_row = summary_table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, "1F4E79")
        cell.paragraphs[0].runs[0].font.color.rgb = None
        # 흰색 글자
        from docx.shared import RGBColor
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    for i, screen in enumerate(SCREENS):
        row = summary_table.rows[i+1]
        row.cells[0].text = screen['id']
        row.cells[1].text = screen['name']
        row.cells[2].text = screen['description'][:50] + "..." if len(screen['description']) > 50 else screen['description']
        row.cells[3].text = screen['access']
    
    doc.add_page_break()
    
    # ========== 2. 화면 상세 ==========
    doc.add_heading("2. 화면 상세", level=1)
    
    for idx, screen in enumerate(SCREENS):
        if idx > 0:
            doc.add_page_break()
        
        # 화면 제목
        doc.add_heading(f"2.{idx+1} {screen['name']} ({screen['id']})", level=2)
        
        # 기본 정보 테이블
        info_table = doc.add_table(rows=3, cols=2)
        info_table.style = 'Table Grid'
        
        info_table.cell(0, 0).text = "화면 ID"
        info_table.cell(0, 1).text = screen['id']
        info_table.cell(1, 0).text = "화면명"
        info_table.cell(1, 1).text = screen['name']
        info_table.cell(2, 0).text = "접근 권한"
        info_table.cell(2, 1).text = screen['access']
        
        for row in info_table.rows:
            row.cells[0].paragraphs[0].runs[0].bold = True
            set_cell_shading(row.cells[0], "E7E6E6")
        
        doc.add_paragraph()
        
        # 화면 설명
        doc.add_heading("화면 설명", level=3)
        doc.add_paragraph(screen['description'])
        
        # 주요 기능
        doc.add_heading("주요 기능/구성 요소", level=3)
        for feature in screen['features']:
            doc.add_paragraph(feature, style='List Bullet')
        
        doc.add_paragraph()
        
        # 화면 캡처
        doc.add_heading("화면 캡처", level=3)
        
        image_path = os.path.join("docs/screenshots", screen['image'])
        if os.path.exists(image_path):
            doc.add_picture(image_path, width=Inches(6.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            doc.add_paragraph(f"[이미지 없음: {screen['image']}]")
    
    # ========== 문서 끝 ==========
    doc.add_page_break()
    doc.add_heading("문서 이력", level=1)
    
    history_table = doc.add_table(rows=2, cols=4)
    history_table.style = 'Table Grid'
    
    history_headers = ["버전", "일자", "작성자", "변경 내용"]
    for i, header in enumerate(history_headers):
        cell = history_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, "1F4E79")
        from docx.shared import RGBColor
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    history_table.rows[1].cells[0].text = "1.0"
    history_table.rows[1].cells[1].text = datetime.now().strftime("%Y-%m-%d")
    history_table.rows[1].cells[2].text = "-"
    history_table.rows[1].cells[3].text = "최초 작성"
    
    # 저장
    output_path = "docs/nCompliance_Screen_Design.docx"
    doc.save(output_path)
    print(f"문서가 저장되었습니다: {output_path}")
    return output_path

if __name__ == "__main__":
    create_document()
